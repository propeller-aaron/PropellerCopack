"""Read and write editable content in Word documents."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from .html_blocks import BLOCK_MARKER, BlockSpec, normalize_compare_text


@dataclass
class DocBlock:
    block_id: str
    label: str
    text: str


def parse_document(path: Path) -> list[DocBlock]:
    document = Document(path)
    blocks: list[DocBlock] = []
    current_id: str | None = None
    current_label = ""
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_id, current_label, current_lines
        if current_id is None:
            return
        blocks.append(
            DocBlock(
                block_id=current_id,
                label=current_label,
                text="\n".join(current_lines).strip(),
            )
        )
        current_id = None
        current_label = ""
        current_lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        marker = BLOCK_MARKER.match(text)
        if marker:
            flush()
            current_id = marker.group(1)
            current_label = ""
            current_lines = []
            continue

        if current_id is None:
            continue

        if not current_label and paragraph.style.name.startswith("Heading"):
            current_label = text
            continue

        if text or current_lines:
            current_lines.append(paragraph.text.rstrip())

        if text == "End of document":
            flush()
            break

    flush()
    return blocks


def write_document(
    path: Path,
    *,
    title: str,
    subtitle: str,
    blocks: list[tuple[BlockSpec, str]],
    notes: list[str] | None = None,
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle_paragraph = document.add_paragraph(subtitle)
    subtitle_paragraph.runs[0].font.color.rgb = RGBColor(0x65, 0x6D, 0x76)

    document.add_paragraph()

    intro = document.add_paragraph(
        "Edit the text under each section. Keep the [[block:...]] lines unchanged. "
        "Use Enter for line breaks (they become <br> on the website). "
        "After saving, run: python tools/apply_content_docs.py"
    )
    intro.runs[0].italic = True

    if notes:
        for note in notes:
            note_paragraph = document.add_paragraph(note)
            note_paragraph.runs[0].italic = True

    document.add_paragraph()

    for spec, text in blocks:
        marker = document.add_paragraph(f"[[block:{spec.block_id}]]")
        marker.runs[0].font.name = "Consolas"
        marker.runs[0].font.size = Pt(9)
        marker.runs[0].font.color.rgb = RGBColor(0x09, 0x69, 0xDA)

        heading = document.add_heading(spec.label, level=3)
        heading.runs[0].font.size = Pt(12)

        body = document.add_paragraph(text.replace("\r\n", "\n").replace("\r", "\n"))
        body.paragraph_format.space_after = Pt(12)

    document.save(path)
